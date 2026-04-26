"""
Shared pytest fixtures.

Creates temporary test files (TXT, MD, PDF, DOCX) used across test modules.
All files are written to a tmp_path directory provided by pytest.
"""

from __future__ import annotations

from pathlib import Path

import pytest


# ── Plain-text / Markdown fixtures ───────────────────────────────────────────

@pytest.fixture()
def txt_file(tmp_path: Path) -> Path:
    p = tmp_path / "sample.txt"
    lines = [f"これはサンプルテキストファイルの {i + 1} 行目です。" for i in range(10)]
    p.write_text("\n".join(lines), encoding="utf-8")
    return p


@pytest.fixture()
def md_file(tmp_path: Path) -> Path:
    p = tmp_path / "sample.md"
    p.write_text(
        "# セクション1\nこれはMarkdownの最初のセクションです。\n\n"
        "## セクション2\nこれは二番目のセクションです。\n\n"
        "## セクション3\n詳細な内容がここに入ります。\n",
        encoding="utf-8",
    )
    return p


# ── PDF fixture ───────────────────────────────────────────────────────────────

@pytest.fixture()
def pdf_file(tmp_path: Path) -> Path:
    """Create a minimal 2-page PDF using fpdf2."""
    try:
        from fpdf import FPDF
    except ImportError:
        pytest.skip("fpdf2 not installed")

    pdf = FPDF()
    for page_num in range(1, 3):
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(0, 10, f"Page {page_num}: This is test content for DocuMate.")
        pdf.ln()
        pdf.cell(0, 10, "RAG stands for Retrieval-Augmented Generation.")
    p = tmp_path / "sample.pdf"
    pdf.output(str(p))
    return p


# ── DOCX fixture ──────────────────────────────────────────────────────────────

@pytest.fixture()
def docx_file(tmp_path: Path) -> Path:
    """Create a minimal DOCX with 5 paragraphs using python-docx."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("テストドキュメント", level=1)
    for i in range(1, 6):
        doc.add_paragraph(f"これは {i} 番目のパラグラフです。社内ドキュメントのテスト用テキストです。")
    p = tmp_path / "sample.docx"
    doc.save(str(p))
    return p
